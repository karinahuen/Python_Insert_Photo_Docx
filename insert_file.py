import time
timestr = time.strftime("%Y%m%d_%H%M%S")
file_name = "Photo " + timestr

from tkinter import filedialog
from tkinter import *
root = Tk()
root.withdraw()
folder_selected = filedialog.askdirectory()

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ctypes import alignment
from docx.enum.text import WD_COLOR_INDEX
 
trow = 0
tcol = 2

document = Document()
table = document.add_table(rows=trow, cols=tcol)
#------Loads Photo Path----------
import cv2
import os
import numpy as np
images = []

#--Table Header--
header_column_cells = table.add_row().cells
header_img = header_column_cells[0].paragraphs[0]
header_img_run = header_img.add_run()
header_img_run.text = "Photo"
header_img_run.font.bold = True
header_img_run.font.underline = True
header_img_run = header_img.add_run()

hedaer_img_desc = header_column_cells[1].paragraphs[0]
hedaer_img_desc_run = hedaer_img_desc.add_run()
hedaer_img_desc_run.text = "Description"
hedaer_img_desc_run.font.bold = True
hedaer_img_desc_run.font.underline = True
hedaer_img_desc_run = hedaer_img_desc.add_run()

for filename in os.listdir(folder_selected):
    #img = cv2.imread(os.path.join(folder_selected,filename)) 
    #--image with unicode characters--
    img = cv2.imdecode(np.fromfile(os.path.join(folder_selected,filename), dtype=np.uint8), cv2.IMREAD_UNCHANGED)
    if img is not None:
        images.append(img)
        pic = folder_selected + "//" + filename

        #--add new line--
        row_cells_des = table.add_row().cells 
        paragraph = row_cells_des[0].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(pic, height = 3900000)
        paragraph_column = row_cells_des[1].paragraphs[0]
        runns = paragraph_column.add_run(filename)
        runns.bold = True
        runns.italic  = True
        row_cells_space = table.add_row().cells 

table.style = 'Table Grid'
document.save(f"{folder_selected}\%s.docx" % file_name)